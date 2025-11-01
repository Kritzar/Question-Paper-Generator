import os
from flask import Flask, render_template, request, send_file, session, redirect, url_for, flash
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_sqlalchemy import SQLAlchemy
import pdfplumber
import docx
import google.generativeai as genai
from fpdf import FPDF
import re
import pdf_extraction
# Configure Gemini API
genai.configure(api_key="xxxxxxxxxxxxxxxxxxxxxx")
model = genai.GenerativeModel("models/gemini-1.5-pro")

# App setup
app = Flask(__name__)
app.secret_key = 'xxxxxxxxxxxxxxxx'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# File settings
app.config['UPLOAD_FOLDER_HOME'] = 'uploads_home/'
app.config['UPLOAD_FOLDER_QUIZ'] = 'uploads_quiz/'
app.config['GENERATED_FOLDER'] = 'generated_questions/'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'docx'}
app.config['MAX_PAGE_LIMIT'] = 50  # Define the maximum page limit

os.makedirs(app.config['UPLOAD_FOLDER_HOME'], exist_ok=True)
os.makedirs(app.config['UPLOAD_FOLDER_QUIZ'], exist_ok=True)
os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)


# User Model
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)


# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def extract_text_from_file(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        with pdfplumber.open(file_path) as pdf:
            return ''.join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif ext == 'docx':
        return ' '.join([p.text for p in docx.Document(file_path).paragraphs])
    elif ext == 'txt':
        with open(file_path, 'r') as file:
            return file.read()
    return None


def get_page_count(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        with pdfplumber.open(file_path) as pdf:
            return len(pdf.pages)
    elif ext == 'docx':
        doc = docx.Document(file_path)
        # This is a rough estimate; DOCX doesn't have a reliable page count.
        # You might adjust this based on your needs (average characters per page, etc.)
        text = " ".join([p.text for p in doc.paragraphs])
        return len(text) // 3000  # Assuming roughly 3000 characters per page
    elif ext == 'txt':
        with open(file_path, 'r') as file:
            text = file.read()
        return len(text) // 3000  # Rough estimate for txt too
    return 0  # Default


def generate_questions(input_text, num_questions, difficulty, question_type):
    prompt = f"""
    You are an AI assistant helping the user generate {question_type} questions based on the following text:
    '{input_text}'
    Please generate {num_questions} {question_type} questions from the text at a {difficulty} difficulty level.

    For Fill in the Blanks, provide the question with underscores and the answer in brackets at the end.
    For MCQs: Question + four options (A, B, C, D), mark correct option like "[C]".
    For Short Answers: Question followed by concise answer in parentheses.

    Separate each question with '##---##'.
    """
    response = model.generate_content(prompt).text.strip()
    return [q.strip() for q in response.split('##---##') if q.strip()]


def generate_quiz_mcqs(context, num_questions, difficulty):
    mcqs = []
    for _ in range(num_questions):
        prompt = f"""Generate multiple-choice questions based on the following text:
        with this level of difficulty {difficulty}
        {context}

        The question should have four options, with one correct answer. Format your response as follows:

        Question: [Your Question Here]
        Options:
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        Answer: [The Correct Option - A, B, C, or D]
        """

        try:
            response = model.generate_content(prompt)
            if response.parts and hasattr(response.parts[0], 'text'):
                mcq_text = response.parts[0].text
                # Parse the generated text to extract question, options, and answer
                question_match = re.search(r"Question:\s*(.+)", mcq_text, re.IGNORECASE)
                options_match = re.search(r"Options:\s*A\)\s*(.+)\s*B\)\s*(.+)\s*C\)\s*(.+)\s*D\)\s*(.+)", mcq_text, re.IGNORECASE | re.DOTALL)
                answer_match = re.search(r"Answer:\s*([A-D])", mcq_text, re.IGNORECASE)

                if question_match and options_match and answer_match:
                    question = question_match.group(1).strip()
                    options = [opt.strip() for opt in options_match.groups()]
                    answer = answer_match.group(1).strip().upper()
                    mcqs.append({"question": question, "options": options, "answer": answer})
                else:
                    print(f"Warning: Could not fully parse MCQ from Gemini response:\n{mcq_text}")
            else:
                print(f"Warning: Empty or unexpected response from Gemini.")
        except Exception as e:
            print(f"Error generating MCQ from Gemini: {e}")
    return mcqs


# Authentication helpers
def is_logged_in():
    return 'user_id' in session


# Routes
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/help')
def help():
    return render_template('help.html')


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if User.query.filter_by(username=username).first():
            return "Username already exists."
        user = User(username=username, password=generate_password_hash(password))
        db.session.add(user)
        db.session.commit()
        return redirect(url_for('login'))
    return render_template('signup.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            return redirect(url_for('index'))
        return "Invalid credentials."
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('user_id', None)
    return redirect(url_for('index'))


@app.route('/generate_test', methods=['POST'])
def generate_combined_test():
    if not is_logged_in():
        return redirect(url_for('login'))

    if 'file' not in request.files:
        flash("No file part", "error")  # Use flash for user feedback
        return redirect(request.url)
    file = request.files['file']

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER_HOME'], filename)
        file.save(file_path)

        page_count = get_page_count(file_path)
        if page_count > app.config['MAX_PAGE_LIMIT']:
            os.remove(file_path)  # Remove the uploaded file
            flash(f"File exceeds the maximum page limit of {app.config['MAX_PAGE_LIMIT']} pages.", "error")
            return redirect(request.url)  # Redirect back to the form

        # text = extract_text_from_file(file_path)
        text,success=pdf_extraction.process_pdf_with_timeout_and_output(file_path,output_file_path)
        astro=pdf_extraction.analyze_physics_text(text)
        if not success:
            flash("Could not extract text from file.", "error")
            return redirect(request.url)
        provided_text=text
        generated_questions = {}

        # Fill in the blanks
        fib_easy_num = int(request.form.get('fib_easy_num_questions', 0))  # Default to 0 if not provided
        fib_hard_num = int(request.form.get('fib_hard_num_questions', 0))
        if fib_easy_num > 0:
            # generated_questions.setdefault('fill_in_the_blanks', []).extend(
            #     generate_questions(text, fib_easy_num, "Easy", "fill in the blank"))
            difficulty_level="easy"
            easy_questions = pdf_extraction.generate_fill_in_the_blanks(text, difficulty_level, fib_easy_num)
            generated_questions.setdefault('fill_in_the_blanks', []).extend(
                [f"{q_a['question']} ({q_a['answer']})" for q_a in easy_questions]
            )
        if fib_hard_num > 0:
            # generated_questions.setdefault('fill_in_the_blanks', []).extend(
            #     generate_questions(text, fib_hard_num, "Hard", "fill in the blank"))
            difficulty_level="hard"
            easy_questions = pdf_extraction.generate_fill_in_the_blanks(text, difficulty_level, fib_hard_num)
            generated_questions.setdefault('fill_in_the_blanks', []).extend(
                [f"{q_a['question']} ({q_a['answer']})" for q_a in easy_questions]
            )

        # Numerical
        numerical_easy_num = int(request.form.get('mcq_easy_num_questions', 0))
        # mcq_medium_num = int(request.form.get('mcq_medium_num_questions', 0))
        numerical_hard_num = int(request.form.get('mcq_hard_num_questions', 0))
        if numerical_easy_num > 0:
            easy_questions = pdf_extraction.generate_n_easy_problems(astro, numerical_easy_num)
            generated_questions.setdefault('numericals', []).extend(
                [f"{q_a['problem']} {q_a['solution']} {q_a['answer']}" for q_a in easy_questions]
            )
            # generated_questions.setdefault('mcqs', []).extend(
            #     generate_questions(text, numerical_easy_num))
        # if mcq_medium_num > 0:
        #     generated_questions.setdefault('mcqs', []).extend(
        #         generate_questions(text, mcq_medium_num, "Medium", "multiple choice"))
        if numerical_hard_num > 0:
            easy_questions = pdf_extraction.generate_n_hard_problems(astro, numerical_hard_num)
            generated_questions.setdefault('numericals', []).extend(
                [f"{q_a['problem']} {q_a['solution']} {q_a['answer']}" for q_a in easy_questions]
            )
            # generated_questions.setdefault('mcqs', []).extend(
            #     generate_questions(text, numerical_hard_num))

        # Short Answers
        short_answer_easy_num = int(request.form.get('short_answer_easy_questions', 0))
        short_answer_medium_num = int(request.form.get('short_answer_medium_questions', 0))
        short_answer_hard_num = int(request.form.get('short_answer_hard_questions', 0))
        generated_questions = {}

        if short_answer_easy_num > 0:
            difficulty_level="easy"
            easy_questions = pdf_extraction.generate_theoretical_qa(text, difficulty_level, short_answer_easy_num)
            generated_questions.setdefault('short_answers', []).extend(
                [f"{q_a['question']} ({q_a['answer']})" for q_a in easy_questions]
            )

        # if short_answer_medium_num > 0:
        #     difficulty_level="easy"
        #     medium_questions = pdf_extraction.generate_theoretical_qa(text, "medium", short_answer_medium_num)
        #     generated_questions.setdefault('short_answers', []).extend(
        #         [f"{q_a['question']} ({q_a['answer']})" for q_a in medium_questions]
        #     )

        if short_answer_hard_num > 0:
            difficulty_level="hard"
            hard_questions = pdf_extraction.generate_theoretical_qa(text, difficulty_level, short_answer_hard_num)
            generated_questions.setdefault('short_answers', []).extend(
                [f"{q_a['question']} ({q_a['answer']})" for q_a in hard_questions]
            )
        if not generated_questions:
            flash("No questions generated.", "info")
            return redirect(request.url)

        all_questions = []
        for qlist in generated_questions.values():
            all_questions.extend(qlist)
        download_data = "|||".join(all_questions)

        return render_template('generated_test.html', generated_questions=generated_questions,
                               download_data=download_data)
    return "Invalid file"


@app.route('/download/<file_type>')
def download_combined_file(file_type):
    data = request.args.get('data')
    if not data:
        return "No data to download."

    questions = data.split("|||")

    if file_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for q in questions:
            pdf.multi_cell(0, 10, q)
            pdf.ln()
        path = os.path.join(app.config['GENERATED_FOLDER'], 'generated_questions.pdf')
        pdf.output(path)
        return send_file(path, as_attachment=True, mimetype='application/pdf',
                        download_name='generated_questions.pdf')

    elif file_type == 'txt':
        path = os.path.join(app.config['GENERATED_FOLDER'], 'generated_questions.txt')
        with open(path, 'w') as f:
            f.write("\n\n".join(questions))
        return send_file(path, as_attachment=True, mimetype='text/plain',
                        download_name='generated_questions.txt')

    return "Invalid format."


@app.route('/quiz')
def quiz_form():
    return render_template('quiz.html')


@app.route('/generate_mcqs', methods=['POST'])
def generate_quiz():
    if not is_logged_in():
        return redirect(url_for('login'))

    if 'file' not in request.files:
        flash("No file selected", "error")
        return redirect(request.url)
    file = request.files['file']

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        path = os.path.join(app.config['UPLOAD_FOLDER_QUIZ'], filename)
        file.save(path)

        page_count = get_page_count(path)
        if page_count > app.config['MAX_PAGE_LIMIT']:
            os.remove(path)
            flash(f"File exceeds the page limit of {app.config['MAX_PAGE_LIMIT']} pages", "error")
            return redirect(request.url)

        text,success=pdf_extraction.process_pdf_with_timeout_and_output(path,output_file_path2)
        
        if not success:
            flash("Could not extract text from file.", "error")
            return redirect(request.url)

        num = int(request.form['num_questions'])
        diff = request.form['difficulty']
        mcqs = generate_quiz_mcqs(text, num, diff)
        session['quiz_questions'] = mcqs
        session['score'] = 0
        session['current_question'] = 0
        return render_template('quiz_attempt.html', question=mcqs[0], question_number=1,
                               total_questions=len(mcqs))
    return "Invalid format"


@app.route('/submit_answer', methods=['POST'])
def submit_quiz_answer():
    if 'quiz_questions' not in session:
        return "Session expired"

    questions = session['quiz_questions']
    idx = session['current_question']
    user_ans = request.form.get('answer')
    if user_ans == questions[idx]['correct_answer']:
        session['score'] += 1
    session['current_question'] += 1

    if session['current_question'] < len(questions):
        return render_template('quiz_attempt.html',
                               question=questions[session['current_question']],
                               question_number=session['current_question'] + 1,
                               total_questions=len(questions))
    else:
        score = session['score']
        total = len(questions)
        session.pop('quiz_questions', None)
        session.pop('score', None)
        session.pop('current_question', None)
        return render_template('quiz_results.html', score=score, total_questions=total, all_questions=questions)


if __name__ == "__main__":
    output_file_path = "final_output.txt"
    output_file_path2= "final_output2.txt"
    with app.app_context():
        db.create_all()

    app.run(debug=True)
